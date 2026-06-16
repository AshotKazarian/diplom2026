from io import BytesIO

import pandas as pd
from django.utils import timezone
from docx.shared import Mm

from accounts.models import Student

from .models import ThesisApplication


def export_registry_to_excel():
    """Экспорт реестра в Excel"""
    try:
        applications = ThesisApplication.objects.select_related("student", "professor")

        data = []
        for app in applications:
            data.append(
                {
                    "ФИО студента": app.student.full_name,
                    "Группа": app.student.group,
                    "Тема ВКР": app.topic,
                    "Научный руководитель": app.professor.full_name,
                    "Кафедра": (
                        app.professor.department.name
                        if app.professor.department
                        else ""
                    ),
                    "Статус": app.get_status_display(),
                    "Дата создания": app.created_at.strftime("%d.%m.%Y"),
                }
            )

        df = pd.DataFrame(data)
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Реестр ВКР", index=False)
        output.seek(0)
        return output
    except Exception as e:
        print(f"Error exporting to Excel: {e}")
        output = BytesIO()
        df = pd.DataFrame()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Реестр ВКР")
        output.seek(0)
        return output


def get_group_info(group_name):
    """Получает информацию о группе"""
    students = Student.objects.filter(group=group_name, user__is_active=True)
    if not students.exists():
        raise ValueError(f"Группа {group_name} не найдена")

    first_student = students.first()
    faculty = first_student.faculty
    department = first_student.department
    profile = first_student.profile

    education_code = (
        profile.education_code if profile and profile.education_code else ""
    )
    education_specialty = (
        profile.education_specialty if profile and profile.education_specialty else ""
    )

    return {
        "faculty": faculty,
        "department": department,
        "education_code": education_code,
        "education_specialty": education_specialty,
        "students": students,
    }


def generate_general_appointment_order(group_name, request):
    """Формирование приказа со стилями из шаблона"""
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        group_info = get_group_info(group_name)
        faculty = group_info["faculty"]
        department = group_info["department"]
        education_code = group_info["education_code"]
        education_specialty = group_info["education_specialty"]

        applications = (
            ThesisApplication.objects.filter(
                status="approved", student__group=group_name
            )
            .select_related("student", "professor")
            .order_by("student__full_name")
        )

        if not applications.exists():
            raise ValueError(f"Нет одобренных заявок для группы {group_name}")

        current_datetime = timezone.now()
        order_date = current_datetime.strftime("%d.%m.%Y")

        doc = Document()

        # Настройка полей страницы
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.87)
            section.bottom_margin = Cm(2.47)
            section.left_margin = Cm(2.75)
            section.right_margin = Cm(1.0)
            section.header_distance = Cm(0)
            section.footer_distance = Cm(2.12)
            section.page_width = Mm(210)  # 21,0 см (210 мм)
            section.page_height = Mm(297)  # 29,7 см (297 мм)

        # ========== СТРАНИЦА 1 ==========

        p1 = doc.add_paragraph()
        run1 = p1.add_run(
            "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ"
        )
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(9)
        run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(3.2)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1
        p1.paragraph_format.left_indent = Cm(0.07)

        p2 = doc.add_paragraph()
        run2 = p2.add_run(
            "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ"
        )
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(8)
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(6.65)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1
        p2.paragraph_format.left_indent = Cm(0.07)
        p2.paragraph_format.right_indent = Cm(0.01)

        p3 = doc.add_paragraph()
        run3 = p3.add_run("«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»")
        run3.font.name = "Times New Roman"
        run3.font.size = Pt(11)
        run3.font.bold = True
        run3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(5.9)
        p3.paragraph_format.space_after = Pt(0)
        p3.paragraph_format.line_spacing = 1
        p3.paragraph_format.left_indent = Cm(0.07)
        p3.paragraph_format.right_indent = Cm(0.05)

        p4 = doc.add_paragraph()
        run4 = p4.add_run("(МОСКОВСКИЙ ПОЛИТЕХ)")
        run4.font.name = "Times New Roman"
        run4.font.size = Pt(14)
        run4.font.bold = True
        run4._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(4.55)
        p4.paragraph_format.space_after = Pt(0)
        p4.paragraph_format.line_spacing = 1
        p4.paragraph_format.left_indent = Cm(0.07)
        p4.paragraph_format.right_indent = Cm(0.01)

        p5 = doc.add_paragraph()
        run5 = p5.add_run("ПРИКАЗ")
        run5.font.name = "Times New Roman"
        run5.font.size = Pt(21)
        run5.font.bold = True
        run5._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.paragraph_format.space_before = Pt(0)
        p5.paragraph_format.space_after = Pt(0)
        p5.paragraph_format.line_spacing = 1

        doc.add_paragraph()

        # Дата и номер приказа (таблица без границ)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(10)
        table.columns[1].width = Cm(10)

        left_cell = table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        run_left = left_para.add_run(order_date)
        run_left.font.name = "Times New Roman"
        run_left.font.size = Pt(14)
        run_left.underline = True
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_para.paragraph_format.space_before = Pt(6.45)
        left_para.paragraph_format.space_after = Pt(0)
        left_para.paragraph_format.line_spacing = 1
        left_para.paragraph_format.left_indent = Cm(0.14)

        right_cell = table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        run_right_prefix = right_para.add_run("№ ")
        run_right_prefix.font.name = "Times New Roman"
        run_right_prefix.font.size = Pt(14)
        run_right_number = right_para.add_run("_____-__")
        run_right_number.font.name = "Times New Roman"
        run_right_number.font.size = Pt(14)
        run_right_number.underline = True
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_para.paragraph_format.space_before = Pt(4.8)
        right_para.paragraph_format.space_after = Pt(0)
        right_para.paragraph_format.line_spacing = 1
        right_para.paragraph_format.left_indent = Cm(0.14)

        for cell in table.rows[0].cells:
            cell._element.get_or_add_tcPr().append(OxmlElement("w:noBorders"))

        doc.add_paragraph()

        p9 = doc.add_paragraph()
        run9 = p9.add_run("О назначении руководителей")
        run9.font.name = "Times New Roman"
        run9.font.size = Pt(14)
        p9.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p9.paragraph_format.space_before = Pt(0.9)
        p9.paragraph_format.space_after = Pt(0)
        p9.paragraph_format.line_spacing = 1
        p9.paragraph_format.left_indent = Cm(0.197)

        p10 = doc.add_paragraph()
        run10 = p10.add_run("и утверждении тем выпускных")
        run10.font.name = "Times New Roman"
        run10.font.size = Pt(14)
        p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p10.paragraph_format.space_before = Pt(0.35)
        p10.paragraph_format.space_after = Pt(0)
        p10.paragraph_format.line_spacing = 1.02
        p10.paragraph_format.left_indent = Cm(0.197)
        p10.paragraph_format.right_indent = Cm(3.593)

        p11 = doc.add_paragraph()
        run11 = p11.add_run("квалификационных работ")
        run11.font.name = "Times New Roman"
        run11.font.size = Pt(14)
        p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p11.paragraph_format.space_before = Pt(0)
        p11.paragraph_format.space_after = Pt(0)
        p11.paragraph_format.line_spacing = 1
        p11.paragraph_format.left_indent = Cm(0.197)

        doc.add_paragraph()
        doc.add_paragraph()

        if education_code and education_specialty:
            education_text = f"В соответствии с федеральным государственным образовательным стандартом высшего образования по направлению подготовки {education_code} «{education_specialty}»"
        else:
            education_text = f"В соответствии с федеральным государственным образовательным стандартом высшего образования"

        p14 = doc.add_paragraph()
        run14 = p14.add_run(education_text)
        run14.font.name = "Times New Roman"
        run14.font.size = Pt(14)
        p14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p14.paragraph_format.space_before = Pt(0.05)
        p14.paragraph_format.space_after = Pt(0)
        p14.paragraph_format.line_spacing = 1.02
        p14.paragraph_format.left_indent = Cm(0.197)
        p14.paragraph_format.right_indent = Cm(0.173)
        p14.paragraph_format.first_line_indent = Cm(0.626)

        p15 = doc.add_paragraph()
        run15 = p15.add_run("ПРИКАЗЫВАЮ:")
        run15.font.name = "Times New Roman"
        run15.font.size = Pt(14)
        run15.font.bold = True
        p15.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p15.paragraph_format.space_before = Pt(2.75)
        p15.paragraph_format.space_after = Pt(0)
        p15.paragraph_format.line_spacing = 1
        p15.paragraph_format.left_indent = Cm(0.834)

        department_name = department.name if department else ""
        faculty_name = faculty.name if faculty else ""

        # Пункт 1 (с отступом, без сложной нумерации)
        p16 = doc.add_paragraph()
        run16 = p16.add_run(
            f"1. Назначить руководителей и утвердить темы выпускных квалификационных работ по кафедре «{department_name}» факультета {faculty_name} следующим обучающимся:"
        )
        run16.font.name = "Times New Roman"
        run16.font.size = Pt(14)
        p16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p16.paragraph_format.space_before = Pt(2.4)
        p16.paragraph_format.space_after = Pt(0)
        p16.paragraph_format.line_spacing = 1.02
        p16.paragraph_format.left_indent = Cm(0)
        p16.paragraph_format.first_line_indent = Cm(0.626)

        p17 = doc.add_paragraph()
        run17 = p17.add_run(f"учебная группа {group_name}:")
        run17.font.name = "Times New Roman"
        run17.font.size = Pt(14)
        p17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p17.paragraph_format.space_before = Pt(2.35)
        p17.paragraph_format.space_after = Pt(1.65)
        p17.paragraph_format.line_spacing = 1
        p17.paragraph_format.left_indent = Cm(0.829)

        # ========== ОСНОВНАЯ ТАБЛИЦА ==========
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Устанавливаем общую ширину таблицы
        # table.width = Cm(14.11)
        
        # Устанавливаем ширину столбцов
        col_widths = [1.11, 4.63, 5.19, 5.19]
        for i, width in enumerate(col_widths):
            table.columns[i].width = Cm(width)
            
            # Дополнительно фиксируем ширину через XML
            for cell in table.columns[i].cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width * 567)))  # см → dxa
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcPr.append(OxmlElement("w:noMar"))

        hdr_cells = table.rows[0].cells
        headers = [
            "№ п/п",
            "Фамилия Имя Отчество",
            "Тема выпускной квалификационной работы",
            "Руководитель (должность, ФИО)",
        ]

        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

        for idx, app in enumerate(applications, 1):
            row_cells = table.add_row().cells

            row_cells[0].text = f"{idx}."
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in row_cells[0].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

            row_cells[1].text = app.student.full_name
            for run in row_cells[1].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

            row_cells[2].text = app.topic
            for run in row_cells[2].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

            professor = app.professor
            professor_title = (
                f"{professor.position.name} " if professor.position else ""
            )
            professor_info = f"{professor.full_name}, {professor_title}"
            row_cells[3].text = professor_info
            for run in row_cells[3].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Основание
        p19 = doc.add_paragraph()
        run19 = p19.add_run(
            f"Основание: протокол заседания кафедры «{department_name}» от __.__.____ № ___/__-__"
        )
        run19.font.name = "Times New Roman"
        run19.font.size = Pt(14)
        p19.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p19.paragraph_format.space_before = Pt(0.5)
        p19.paragraph_format.space_after = Pt(0)
        p19.paragraph_format.line_spacing = 1.02
        p19.paragraph_format.left_indent = Cm(0.197)
        p19.paragraph_format.right_indent = Cm(0.176)
        p19.paragraph_format.first_line_indent = Cm(0.626)

        # Пункт 2
        p20 = doc.add_paragraph()
        run20 = p20.add_run(
            "2. Контроль за исполнением приказа возложить на соответствующего заведующего кафедрой и/или руководителя образовательной программы."
        )
        run20.font.name = "Times New Roman"
        run20.font.size = Pt(14)
        p20.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p20.paragraph_format.space_before = Pt(0)
        p20.paragraph_format.space_after = Pt(0)
        p20.paragraph_format.line_spacing = 1.02
        p20.paragraph_format.left_indent = Cm(0)
        p20.paragraph_format.right_indent = Cm(0.173)
        p20.paragraph_format.first_line_indent = Cm(0.626)

        # Отступ между пунктами
        doc.add_paragraph()

        # Подпись (Декан)
        p23 = doc.add_paragraph()
        run23 = p23.add_run("Декан факультета")
        run23.font.name = "Times New Roman"
        run23.font.size = Pt(14)
        p23.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p23.paragraph_format.space_before = Pt(2.45)
        p23.paragraph_format.space_after = Pt(0)
        p23.paragraph_format.line_spacing = 1
        p23.paragraph_format.left_indent = Cm(0.197)

        # ========== НИЖНИЙ КОЛОНТИТУЛ ==========
        footer = doc.sections[0].footer
        footer.paragraphs[0].clear()

        footer_para1 = footer.add_paragraph()
        run_f1 = footer_para1.add_run(
            "О назначении руководителей и утверждении тем выпускных квалификационных работ – 09-21"
        )
        run_f1.font.name = "Times New Roman"
        run_f1.font.size = Pt(9)
        footer_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para1.paragraph_format.space_before = Pt(0.6)
        footer_para1.paragraph_format.space_after = Pt(0)
        footer_para1.paragraph_format.line_spacing = 1.02
        footer_para1.paragraph_format.left_indent = Cm(0.02)

        footer_para2 = footer.add_paragraph()
        run_f2 = footer_para2.add_run(
            f"Центр по работе со студентами, отделение «На Большой Семеновской» (очная форма обучения) Факультет {faculty_name}"
        )
        run_f2.font.name = "Times New Roman"
        run_f2.font.size = Pt(9)
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para2.paragraph_format.space_before = Pt(0)
        footer_para2.paragraph_format.space_after = Pt(0)
        footer_para2.paragraph_format.line_spacing = 1.02
        footer_para2.paragraph_format.left_indent = Cm(0.02)

        footer_para3 = footer.add_paragraph()
        run_f3 = footer_para3.add_run(
            f"Исп.: Фамилия Имя Отчество; тел.: +7 (___) ___-__-__, дата формирования: {order_date}"
        )
        run_f3.font.name = "Times New Roman"
        run_f3.font.size = Pt(9)
        footer_para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para3.paragraph_format.space_before = Pt(0)
        footer_para3.paragraph_format.space_after = Pt(0)
        footer_para3.paragraph_format.line_spacing = 1
        footer_para3.paragraph_format.left_indent = Cm(0.02)

        footer_para4 = footer.add_paragraph()
        run_f4 = footer_para4.add_run("ИД ")
        run_f4.font.name = "Times New Roman"
        run_f4.font.size = Pt(9)
        footer_para4.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para4.paragraph_format.space_before = Pt(0.25)
        footer_para4.paragraph_format.space_after = Pt(0)
        footer_para4.paragraph_format.line_spacing = 1
        footer_para4.paragraph_format.left_indent = Cm(0.02)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"Error generating document: {e}")
        import traceback

        traceback.print_exc()
        return None
